from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _set_run_font_simsun(run):
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def _set_paragraph_font_simsun(paragraph):
    for run in paragraph.runs:
        _set_run_font_simsun(run)


def build_docx(output_path: Path) -> None:
    doc = Document()

    # Page setup: A4 portrait
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    title = doc.add_paragraph("天主教、基督教、东正教、犹太教、伊斯兰教：相同点与不同点对比")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_font_simsun(title)
    for run in title.runs:
        run.font.size = Pt(16)
        run.bold = True

    doc.add_paragraph("")  # spacer

    p = doc.add_paragraph("相同点（五大一神教的共同特征）")
    _set_paragraph_font_simsun(p)
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(12)

    same_points = [
        "都强调唯一真神/至高者（一神信仰传统）。",
        "都有经典与传统解释体系（经文、注释、教法/教规/教义）。",
        "重视祈祷/礼拜与伦理（行善、悔改、节制、慈善等）。",
        "都有共同体组织（会堂/教会/清真寺与宗教领袖）。",
        "普遍相信末世审判与来世（但细节不同）。",
    ]
    for s in same_points:
        bp = doc.add_paragraph(style=None)
        bp.style = doc.styles["List Bullet"]
        run = bp.add_run(s)
        _set_run_font_simsun(run)
        run.font.size = Pt(11)

    doc.add_page_break()

    h = doc.add_paragraph("不同点对比表")
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_font_simsun(h)
    for run in h.runs:
        run.bold = True
        run.font.size = Pt(12)

    headers = [
        "比较维度",
        "天主教",
        "基督教（常指新教）",
        "东正教",
        "犹太教",
        "伊斯兰教",
    ]

    rows = [
        (
            "起源与核心人物",
            "源自耶稣与宗徒传统；罗马主教（教宗）为普世领袖",
            "源自耶稣；16世纪宗教改革后形成多教派",
            "源自早期大公教会传统；以各自治教会为主",
            "起源于古以色列民族信仰；以摩西之约为核心",
            "起源于穆罕默德（7世纪）；自认为延续亚伯拉罕传统",
        ),
        (
            "对“上帝/真主”的理解",
            "一神；三位一体（圣父、圣子、圣灵）",
            "一神；三位一体（多数教派）",
            "一神；三位一体",
            "严格一神（拒绝三位一体）",
            "严格一神（真主独一；拒绝三位一体）",
        ),
        (
            "对耶稣的定位",
            "耶稣为神子、救主；受死复活",
            "耶稣为救主；强调因信称义（教派细节不同）",
            "耶稣为神子、救主；强调神化（theosis）传统",
            "多被视为历史教师/先知之一（传统上不承认为弥赛亚/神子）",
            "耶稣（尔撒）是重要先知、弥赛亚之一；非神子；对受难/复活的理解与基督宗教不同",
        ),
        (
            "经典（主要）",
            "《圣经》（含次经/第二正典）；传统与教会训导权重要",
            "《圣经》（多为66卷）；强调“唯独圣经”（总体）",
            "《圣经》＋圣传/礼仪传统并重",
            "《塔纳赫》（希伯来圣经）＋《塔木德》等拉比传统",
            "《古兰经》＋圣训（Hadith）等",
        ),
        (
            "权威与组织结构",
            "层级制明显；教宗—主教—司铎—执事",
            "多元：长老制/会众制/主教制等；无统一最高领袖",
            "以主教制与各自治教会为主；无单一全球最高领袖（普世牧首具荣誉地位）",
            "会堂与拉比体系；无全球统一神职层级",
            "清真寺与伊玛目/学者；逊尼派较分散、什叶派部分体系更集中",
        ),
        (
            "救赎/得救观",
            "信德＋圣礼＋善工（在恩宠中）",
            "普遍强调“因信称义”；善行作为信心结果（教派不同）",
            "强调在恩典中与神联合、成圣（神化）",
            "强调遵行律法与盟约生活、悔改（来世观多样）",
            "信仰真主与行善、遵守五功等；强调审判与慈悯",
        ),
        (
            "重要仪式/礼拜",
            "弥撒、圣体、告解等；礼仪统一性较强",
            "主日礼拜、讲道、圣餐/洗礼（视教派）",
            "圣礼仪高度仪式化；圣像与礼仪传统突出",
            "安息日礼拜、逾越节等节期仪式；割礼等",
            "五次礼拜、聚礼（主麻）、斋月封斋、朝觐等",
        ),
        (
            "圣礼/仪式数量",
            "七件圣事（洗礼、坚振、圣体、告解、傅油、圣秩、婚配）",
            "多承认两大圣礼（洗礼、圣餐；名称与理解不同）",
            "七件圣事（称“神圣奥秘”）",
            "无“圣礼”体系；以诫命/律法规定的仪式为核心",
            "无“圣礼”体系；以五功与教法规训为核心",
        ),
        (
            "圣职与婚姻",
            "拉丁礼司铎多守独身；修会制度发达",
            "多数教派牧师可结婚；修会制度较少或不同",
            "司铎可婚（多限婚前）；主教多由修士出身",
            "拉比可婚；强调家庭生活",
            "伊玛目可婚；允许一夫多妻（有限制，地区与法律不同）",
        ),
        (
            "宗教法/规范",
            "教会法；伦理与教导体系完整",
            "强调信仰与圣经伦理；宗派规范不同",
            "教规与传统并重",
            "哈拉卡（犹太法）体系发达",
            "沙里亚（伊斯兰法）体系发达",
        ),
        (
            "主要节日",
            "圣诞、复活节、圣灵降临等",
            "圣诞、复活节等（各派实践不同）",
            "复活节（帕斯哈）极核心；圣诞等",
            "逾越节、赎罪日、住棚节、光明节等",
            "开斋节、宰牲节等（各地不一）",
        ),
        (
            "象征与场所",
            "十字架；教堂/大教堂",
            "十字架；教堂",
            "十字架＋圣像；教堂",
            "六芒星等；会堂",
            "新月（文化象征）；清真寺",
        ),
        (
            "语言与传播",
            "全球化；拉丁礼传统（现多用本地语言）",
            "全球化；多用本地语言",
            "历史上重礼仪语言与传统（现亦本地化）",
            "与希伯来语传统关联紧密；族群与宗教双重身份常见",
            "全球化；阿拉伯语在礼拜与经典中核心地位",
        ),
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        _set_run_font_simsun(run)
        run.bold = True
        run.font.size = Pt(11)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r in rows:
        row_cells = table.add_row().cells
        for i, text in enumerate(r):
            cell = row_cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            _set_run_font_simsun(run)
            run.font.size = Pt(10.5)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


if __name__ == "__main__":
    out = Path(__file__).resolve().parent / "宗教对比表.docx"
    build_docx(out)
    print(f"Wrote: {out}")

